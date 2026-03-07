from __future__ import annotations

import argparse
from pathlib import Path

import pandas as pd

from scripts.generate_figures import (
    plot_evppi,
    plot_policy_bars,
    plot_uncertainty_decomposition,
)
from scripts.reporting_common import build_reporting_bundle, write_reporting_tables

try:
    from docx import Document
    from docx.shared import Inches
except Exception:  # pragma: no cover - optional dependency surface
    Document = None
    Inches = None

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.units import inch
    from reportlab.pdfgen import canvas
except Exception:  # pragma: no cover - optional dependency surface
    letter = None
    inch = None
    canvas = None


def _format_currency(value: float) -> str:
    return f"${value:,.0f}"


def _format_interval(lo: float, hi: float) -> str:
    return f"{_format_currency(lo)} to {_format_currency(hi)}"


def _describe_jurisdiction(summary: pd.DataFrame) -> tuple[str, str]:
    ordered = summary.sort_values("nb_mean", ascending=False).reset_index(drop=True)
    leader = ordered.iloc[0]
    lead_sentence = (
        f"{leader['policy'].replace('_', ' ').title()} leads on mean net benefit at "
        f"{_format_currency(float(leader['nb_mean']))} "
        f"(90% interval {_format_interval(float(leader['nb_p05']), float(leader['nb_p95']))})."
    )

    if len(ordered) > 1:
        runner_up = ordered.iloc[1]
        gap = float(leader["nb_mean"]) - float(runner_up["nb_mean"])
        gap_sentence = (
            f"The margin over {runner_up['policy'].replace('_', ' ').title()} is "
            f"{_format_currency(gap)} on mean net benefit."
        )
    else:
        gap_sentence = "Only one policy surface is available in this reporting bundle."

    return lead_sentence, gap_sentence


def _top_uncertainty_driver(evppi: pd.DataFrame) -> str | None:
    if evppi.empty:
        return None
    top = evppi.sort_values("evppi", ascending=False).iloc[0]
    return (
        f"Top uncertainty driver: {top['group'].replace('_', ' ')} "
        f"(EVPPI {_format_currency(float(top['evppi']))})."
    )


def _write_docx(
    out_docx: Path,
    title: str,
    sections: list[tuple[str, str]],
    tables: list[tuple[str, pd.DataFrame]],
    figures: list[tuple[str, Path]],
) -> Path:
    if Document is None or Inches is None:
        message = "DOCX output requires python-docx to be installed."
        raise RuntimeError(message)

    doc = Document()
    doc.add_heading(title, level=0)

    for heading, body in sections:
        doc.add_heading(heading, level=1)
        for paragraph in body.split("\n"):
            if paragraph.strip():
                doc.add_paragraph(paragraph.strip())

    for table_title, df in tables:
        doc.add_heading(table_title, level=2)
        table = doc.add_table(rows=1, cols=df.shape[1])
        header_cells = table.rows[0].cells
        for index, column in enumerate(df.columns):
            header_cells[index].text = str(column)
        for _, row in df.iterrows():
            cells = table.add_row().cells
            for index, column in enumerate(df.columns):
                value = row[column]
                cells[index].text = "" if pd.isna(value) else f"{value}"
        doc.add_paragraph()

    for figure_title, figure_path in figures:
        if figure_path.exists():
            doc.add_heading(figure_title, level=2)
            doc.add_picture(str(figure_path), width=Inches(6.5))
            caption_path = figure_path.with_name(f"{figure_path.stem}_caption.md")
            if caption_path.exists():
                doc.add_paragraph(caption_path.read_text(encoding="utf-8").strip())
            doc.add_paragraph()

    out_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_docx))
    return out_docx


def _write_pdf(
    out_pdf: Path,
    title: str,
    sections: list[tuple[str, str]],
    figures: list[tuple[str, Path]],
) -> Path:
    if letter is None or inch is None or canvas is None:
        message = "PDF output requires reportlab to be installed."
        raise RuntimeError(message)

    out_pdf.parent.mkdir(parents=True, exist_ok=True)
    pdf = canvas.Canvas(str(out_pdf), pagesize=letter)
    width, height = letter

    y_position = height - 0.8 * inch
    pdf.setFont("Helvetica-Bold", 16)
    pdf.drawString(0.75 * inch, y_position, title)
    y_position -= 0.5 * inch
    pdf.setFont("Helvetica", 10)

    for heading, body in sections:
        if y_position < 1.5 * inch:
            pdf.showPage()
            y_position = height - 0.8 * inch
            pdf.setFont("Helvetica", 10)

        pdf.setFont("Helvetica-Bold", 12)
        pdf.drawString(0.75 * inch, y_position, heading)
        y_position -= 0.25 * inch
        pdf.setFont("Helvetica", 10)

        for line in body.split("\n"):
            if not line.strip():
                y_position -= 0.12 * inch
                continue
            if y_position < 1.0 * inch:
                pdf.showPage()
                y_position = height - 0.8 * inch
                pdf.setFont("Helvetica", 10)
            pdf.drawString(0.75 * inch, y_position, line[:120])
            y_position -= 0.14 * inch
        y_position -= 0.1 * inch

    for figure_title, figure_path in figures:
        if not figure_path.exists():
            continue
        pdf.showPage()
        y_position = height - 0.8 * inch
        pdf.setFont("Helvetica-Bold", 12)
        pdf.drawString(0.75 * inch, y_position, figure_title)
        y_position -= 0.3 * inch
        pdf.drawImage(
            str(figure_path),
            0.75 * inch,
            0.75 * inch,
            width=width - 1.5 * inch,
            height=height - 1.8 * inch,
            preserveAspectRatio=True,
            anchor="c",
        )
        caption_path = figure_path.with_name(f"{figure_path.stem}_caption.md")
        if caption_path.exists():
            pdf.setFont("Helvetica", 9)
            pdf.drawString(0.75 * inch, 0.45 * inch, caption_path.read_text(encoding="utf-8").strip()[:140])

    pdf.save()
    return out_pdf


def _build_markdown_brief(
    title: str,
    author: str,
    meta_dir: Path,
    bundle: dict[str, object],
) -> str:
    policy_summary = bundle["policy_summary"]
    evppi_by_group = bundle["evppi_by_group"]
    uncertainty_decomposition = bundle["uncertainty_decomposition"]
    manifests = bundle["manifests"]

    lines = [f"# {title}", ""]
    if author:
        lines.append(f"**Author:** {author}")
        lines.append("")
    lines.append("## Executive summary")
    lines.append("")
    lines.append(
        "- This pack summarises policy performance using uncertainty draws and ledger-based net benefit outputs."
    )
    if not policy_summary.empty:
        global_best = policy_summary.sort_values("nb_mean", ascending=False).iloc[0]
        lines.append(
            "- Highest mean net benefit in the current bundle: "
            f"{global_best['policy'].replace('_', ' ').title()} in "
            f"{global_best['jurisdiction'].replace('_', ' ').title()} "
            f"at {_format_currency(float(global_best['nb_mean']))}."
        )
    if not evppi_by_group.empty:
        top_driver = _top_uncertainty_driver(evppi_by_group)
        if top_driver:
            lines.append(f"- {top_driver}")
    lines.append("")

    for jurisdiction in sorted(bundle["run_dirs"]):
        title_case = jurisdiction.replace("_", " ").title()
        jurisdiction_summary = policy_summary.loc[
            policy_summary["jurisdiction"] == jurisdiction
        ].drop(columns=["jurisdiction"])
        lead_sentence, gap_sentence = _describe_jurisdiction(jurisdiction_summary)
        lines.append(f"## {title_case} results")
        lines.append("")
        lines.append(lead_sentence)
        lines.append("")
        lines.append(gap_sentence)
        lines.append("")
        lines.append(jurisdiction_summary.head(3).to_markdown(index=False))
        lines.append("")
        lines.append(
            "Metric note: `prem_*` columns are premium indices from the insurance model, not quoted retail premiums."
        )
        lines.append("")

        jurisdiction_evppi = evppi_by_group.loc[
            evppi_by_group["jurisdiction"] == jurisdiction
        ].drop(columns=["jurisdiction"])
        if not jurisdiction_evppi.empty:
            lines.append(f"### {title_case}: EVPPI by group")
            lines.append("")
            top_driver = _top_uncertainty_driver(jurisdiction_evppi)
            if top_driver:
                lines.append(top_driver)
                lines.append("")
            lines.append(jurisdiction_evppi.to_markdown(index=False))
            lines.append("")

        jurisdiction_decomposition = uncertainty_decomposition.loc[
            uncertainty_decomposition["jurisdiction"] == jurisdiction
        ].drop(columns=["jurisdiction"])
        if not jurisdiction_decomposition.empty:
            lines.append(f"### {title_case}: uncertainty decomposition")
            lines.append("")
            lines.append(jurisdiction_decomposition.to_markdown(index=False))
            lines.append("")

    lines.append("## Reproducibility")
    lines.append("")
    lines.append(
        "This analysis is configuration-driven and each source run carries a manifest plus deterministic reporting outputs."
    )
    lines.append("")

    for jurisdiction, manifest in manifests.items():
        if not manifest:
            continue
        run_id = bundle["run_dirs"][jurisdiction].name
        lines.append(f"### {jurisdiction.replace('_', ' ').title()} run manifest (excerpt)")
        lines.append("")
        lines.append(f"- run_id: {run_id}")
        for key in [
            "created_utc",
            "repo_tree_hash",
            "policies_file",
            "policies_file_sha256",
            "base_config_file_sha256",
        ]:
            value = manifest.get(key)
            if value is not None:
                lines.append(f"- {key}: {value}")
        lines.append("")

    return "\n".join(lines).rstrip() + "\n"


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--meta_dir",
        required=True,
        help="Path to outputs/runs/meta_pipeline/<timestamp>",
    )
    parser.add_argument(
        "--title",
        default="Policy brief: genetic discrimination protections (Australia and New Zealand)",
    )
    parser.add_argument("--author", default="")
    parser.add_argument(
        "--out",
        default="",
        help="Output folder (default: <meta_dir>/publish_pack)",
    )
    args = parser.parse_args()

    meta_dir = Path(args.meta_dir)
    if not meta_dir.exists():
        raise FileNotFoundError(meta_dir)

    out_dir = Path(args.out) if args.out else (meta_dir / "publish_pack")
    out_dir.mkdir(parents=True, exist_ok=True)

    bundle = build_reporting_bundle(meta_dir=meta_dir)
    generated_tables = write_reporting_tables(out_dir, bundle)

    figures_dir = out_dir / "figures"
    figures_dir.mkdir(parents=True, exist_ok=True)
    for jurisdiction in sorted(bundle["run_dirs"]):
        plot_policy_bars(bundle["policy_summary"], jurisdiction, figures_dir, 300, ["png"])
        plot_evppi(bundle["evppi_by_group"], jurisdiction, figures_dir, 300, ["png"])
        plot_uncertainty_decomposition(
            bundle["uncertainty_decomposition"],
            jurisdiction,
            figures_dir,
            300,
            ["png"],
        )

    policy_summary = bundle["policy_summary"]
    evppi_by_group = bundle["evppi_by_group"]
    uncertainty_decomposition = bundle["uncertainty_decomposition"]

    (out_dir / "POLICY_BRIEF.md").write_text(
        _build_markdown_brief(args.title, args.author, meta_dir, bundle),
        encoding="utf-8",
    )

    sections = [
        (
            "Executive summary",
            "This pack summarises policy performance using uncertainty draws and ledger-based net benefit outputs.\n"
            "Policies are ranked by mean net benefit with 90% intervals computed from the reporting bundle.\n"
            "Figures and tables are generated directly from the run manifests and uncertainty outputs.",
        ),
        (
            "Reproducibility",
            "Each generated artifact records its source run identifier and manifest excerpt in reporting_manifest.json.",
        ),
    ]

    tables: list[tuple[str, pd.DataFrame]] = []
    figures: list[tuple[str, Path]] = []
    for jurisdiction in sorted(bundle["run_dirs"]):
        label = jurisdiction.replace("_", " ").title()
        jurisdiction_summary = policy_summary.loc[
            policy_summary["jurisdiction"] == jurisdiction
        ].drop(columns=["jurisdiction"])
        lead_sentence, gap_sentence = _describe_jurisdiction(jurisdiction_summary)
        sections.insert(
            len(sections) - 1,
            (
                f"{label} - top policies",
                f"{lead_sentence}\n{gap_sentence}",
            ),
        )
        tables.append((f"{label} policy summary", jurisdiction_summary))

        jurisdiction_evppi = evppi_by_group.loc[
            evppi_by_group["jurisdiction"] == jurisdiction
        ].drop(columns=["jurisdiction"])
        if not jurisdiction_evppi.empty:
            tables.append((f"{label} EVPPI by group", jurisdiction_evppi))

        jurisdiction_decomposition = uncertainty_decomposition.loc[
            uncertainty_decomposition["jurisdiction"] == jurisdiction
        ].drop(columns=["jurisdiction"])
        if not jurisdiction_decomposition.empty:
            tables.append((f"{label} uncertainty decomposition", jurisdiction_decomposition))

        figures.append(
            (
                f"{label}: net benefit by policy",
                figures_dir / f"{jurisdiction}_net_benefit.png",
            )
        )
        figures.append(
            (
                f"{label}: EVPPI by group",
                figures_dir / f"{jurisdiction}_evppi.png",
            )
        )
        if not jurisdiction_decomposition.empty:
            figures.append(
                (
                    f"{label}: uncertainty decomposition",
                    figures_dir / f"{jurisdiction}_uncertainty_decomposition.png",
                )
            )

    generated_docx = _write_docx(
        out_dir / "POLICY_BRIEF.docx", args.title, sections, tables, figures
    )
    generated_pdf = _write_pdf(out_dir / "POLICY_BRIEF.pdf", args.title, sections, figures)

    if not generated_docx.exists() or not generated_pdf.exists():
        message = "Publish pack did not produce the expected DOCX/PDF artifacts."
        raise RuntimeError(message)

    print("Wrote publish pack to:", out_dir)
    print("Generated tables:", ", ".join(sorted(generated_tables)))
    print("Generated documents:", generated_docx.name, generated_pdf.name)


if __name__ == "__main__":
    main()
